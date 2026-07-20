"""Focused multi-experience tests for validation, rendering, and layout."""

from __future__ import annotations

import shutil
import sys
import tempfile
import unittest
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn


SCRIPTS = Path(__file__).resolve().parents[1]
REPO_ROOT = Path(__file__).resolve().parents[5]
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import check  # noqa: E402
import estimate_layout  # noqa: E402
import render  # noqa: E402


LONG = "Built a deterministic fictional service that improved reliability for every test client."


def _employer(company: str, *, direct=True, projects=None):
    return {
        "company": company,
        "role": "Software Engineer",
        "dates": "2020 – Present",
        "location": "Remote (US)",
        "bullets": [LONG] if direct else [],
        "projects": projects or [],
    }


def _resume(employers, summary_count=3):
    return {
        "name": "Jordan Rivers",
        "contact_line": "City, ST • jordan.rivers@example.com",
        "summary_bullets": [LONG] * summary_count,
        "education_line": "B.S. Computer Science, Example University, 2020",
        "skills": [{"label": "Programming Languages", "items": "Python, Go"}],
        "employers": employers,
    }


def _project(index: int):
    return {"title": f"Synthetic Project {index}",
            "bullets": [LONG, LONG.replace("Built", "Designed")]}


class ValidationTests(unittest.TestCase):
    def test_every_employer_count_order_and_identity_are_locked(self):
        baseline = _resume([
            _employer("Example Systems"),
            _employer("Fictional Labs"),
        ])
        reordered = _resume([
            _employer("Fictional Labs"),
            _employer("Example Systems"),
        ])
        c = check.Checker()
        check.check_locked_fields(c, reordered, baseline)
        self.assertEqual(len(c.failures), 2)

        missing = _resume([_employer("Example Systems")])
        c = check.Checker()
        check.check_locked_fields(c, missing, baseline)
        self.assertTrue(any("Employer count changed" in failure for failure in c.failures))

    def test_summary_count_follows_baseline(self):
        baseline = _resume([_employer("Example Systems")], summary_count=2)
        candidate = _resume([_employer("Example Systems")], summary_count=2)
        c = check.Checker()
        check.check_structure(c, candidate, baseline)
        self.assertEqual(c.failures, [])

        candidate["summary_bullets"].append(LONG)
        c = check.Checker()
        check.check_structure(c, candidate, baseline)
        self.assertTrue(any("exactly 2" in failure for failure in c.failures))

    def test_project_only_resume_allows_one_last_resort_reduction(self):
        baseline = _resume([
            _employer("Example Systems", direct=False,
                      projects=[_project(i) for i in range(5)])
        ])
        candidate = _resume([
            _employer("Example Systems", direct=False,
                      projects=[_project(i) for i in range(4)])
        ])
        c = check.Checker()
        check.check_structure(c, candidate, baseline)
        self.assertEqual(c.failures, [])

        candidate["employers"][0]["projects"].pop()
        c = check.Checker()
        check.check_structure(c, candidate, baseline)
        self.assertTrue(any("Total project count 3" in failure for failure in c.failures))

    def test_baseline_direct_bullet_shape_cannot_be_dropped_or_invented(self):
        baseline = _resume([_employer("Example Systems")])
        dropped = _resume([
            _employer("Example Systems", direct=False, projects=[_project(1)])
        ])
        c = check.Checker()
        check.check_structure(c, dropped, baseline)
        self.assertTrue(any("dropped all direct role bullets" in f for f in c.failures))

        project_only_baseline = _resume([
            _employer("Example Systems", direct=False, projects=[_project(1)])
        ])
        invented = _resume([
            _employer("Example Systems", direct=True, projects=[_project(1)])
        ])
        c = check.Checker()
        check.check_structure(c, invented, project_only_baseline)
        self.assertTrue(any("added direct role bullets" in f for f in c.failures))

    def test_empty_employer_is_rejected(self):
        data = _resume([_employer("Example Systems", direct=False)])
        c = check.Checker()
        check.check_structure(c, data, data)
        self.assertTrue(any("must have direct bullets and/or" in f for f in c.failures))

    def test_direct_and_project_bullets_are_checked(self):
        data = _resume([_employer("Example Systems", projects=[_project(1)])])
        data["employers"][0]["bullets"][0] = "too short"
        data["employers"][0]["projects"][0]["bullets"][0] = "also short"
        c = check.Checker()
        check.check_lengths(c, data)
        self.assertEqual(len(c.failures), 2)

    def test_bullet_length_boundaries_are_inclusive(self):
        data = _resume([_employer("Example Systems")])
        for length, expected_failures in ((44, 1), (45, 0), (215, 0), (216, 1)):
            with self.subTest(length=length):
                data["employers"][0]["bullets"] = ["x" * length]
                c = check.Checker()
                check.check_lengths(c, data)
                self.assertEqual(len(c.failures), expected_failures)

    def test_unmatched_bold_markers_fail(self):
        data = _resume([_employer("Example Systems")])
        data["employers"][0]["bullets"][0] += " **broken"
        c = check.Checker()
        check.check_bold_markers(c, data)
        self.assertEqual(len(c.failures), 1)

    def test_profile_project_ownership_prevents_moving_projects(self):
        profile = """## Experience
### Example Systems — Software Engineer (2020 – Present, Remote)
#### [draft] Owned Project
- A real synthetic bullet
### Other Labs — Staff Engineer (2018 – 2020, Remote)
#### [backup] Other Project
- Another synthetic bullet
"""
        data = _resume([
            _employer("Other Labs", direct=False,
                      projects=[{"title": "Owned Project", "bullets": [LONG]}])
        ])
        data["employers"][0]["role"] = "Staff Engineer"
        c = check.Checker()
        check.check_titles(
            c, data, check.parse_profile_titles(profile),
            check.parse_profile_project_owners(profile))
        self.assertTrue(any("does not belong" in failure for failure in c.failures))

    def test_skill_mentions_are_boundary_aware_and_non_negated(self):
        self.assertFalse(check._mentioned_in_jd("Go", "Support ongoing JavaScript work."))
        self.assertFalse(check._mentioned_in_jd("Java", "Build JavaScript applications."))
        self.assertFalse(check._mentioned_in_jd(
            "Kubernetes", "Kubernetes is not required for this position."))
        self.assertTrue(check._mentioned_in_jd("Go", "Production experience with Go is required."))
        self.assertTrue(check._mentioned_in_jd(
            "Kubernetes", "Operate K8s clusters in production."))

    def test_nested_profile_skills_keep_selective_gating(self):
        approved = ["AWS", "Python"]
        weak = ["AWS (Lambda, SQS, SNS)"]
        self.assertTrue(check._in_list("AWS", approved))
        self.assertFalse(check._in_list("AWS Lambda", approved))
        self.assertTrue(check._in_list("Lambda", weak))
        self.assertTrue(check._in_list("AWS Lambda", weak))
        self.assertTrue(check._mentioned_in_jd(
            "AWS Lambda", "Build event handlers with Lambda and Python."))

    def test_role_filename_collisions_and_missing_jds_fail_closed(self):
        with tempfile.TemporaryDirectory() as tmp:
            app = Path(tmp)
            (app / "meta.yaml").write_text(yaml.safe_dump({
                "jobs": [
                    {"role": "Senior Platform Engineer"},
                    {"role": "Senior-Platform Engineer"},
                ],
            }), encoding="utf-8")
            c = check.Checker()
            check.check_role_filename_collisions(c, app)
            self.assertTrue(any("filename collision" in f for f in c.failures))

            fixture_meta = yaml.safe_load((
                REPO_ROOT / "examples" / "fixtures" / "resume-writer"
                / "_test_application_multi-experience" / "application" / "meta.yaml"
            ).read_text())
            (app / "meta.yaml").write_text(
                yaml.safe_dump(fixture_meta), encoding="utf-8")
            c = check.Checker()
            check.check_application_metadata(c, app)
            self.assertTrue(any("jd_file" in f or "does not exist" in f for f in c.failures))


class RenderAndLayoutTests(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())
        self.addCleanup(shutil.rmtree, self.tmp, ignore_errors=True)
        self.reference = REPO_ROOT / "examples" / "templates" / "reference.example.docx"

    def test_render_outputs_every_employer_and_direct_bullet(self):
        data = _resume([
            _employer("Example Systems"),
            _employer("Fictional Labs", projects=[_project(1)]),
        ])
        output = self.tmp / "resume.docx"
        render.render_from_reference(self.reference, data, output)
        doc = Document(output)
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Example Systems – Software Engineer", text)
        self.assertIn("Fictional Labs – Software Engineer", text)
        self.assertEqual(text.count(LONG), 6)  # summaries + direct bullets + project bullet
        self.assertIn("Synthetic Project 1", text)

        employer_paragraphs = [
            p for p in doc.paragraphs
            if " – Software Engineer" in p.text
        ]
        self.assertEqual(len(employer_paragraphs), 2)
        for paragraph in employer_paragraphs:
            tabs = paragraph._p.findall(".//" + qn("w:tab"))
            # One tab-stop declaration in pPr and one tab character in the run.
            self.assertEqual(len(tabs), 2)

    def test_extra_employer_increases_estimate_even_without_extra_bullets(self):
        metrics = estimate_layout.read_template_metrics(self.reference)
        params = estimate_layout.derived_params(metrics)
        one = _resume([_employer("Example Systems")])
        two = _resume([
            _employer("Example Systems"),
            _employer("Fictional Labs", direct=False),
        ])
        a = estimate_layout.estimate(one, params)
        b = estimate_layout.estimate(two, params)
        self.assertAlmostEqual(
            b["total_pt"] - a["total_pt"], params["extra_employer_pt"], places=3)


if __name__ == "__main__":
    unittest.main()

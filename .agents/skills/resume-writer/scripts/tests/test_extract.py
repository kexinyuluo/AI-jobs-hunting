"""Focused tests for paragraph-based DOCX resume extraction."""

from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).resolve().parents[1]
EXTRACT_SCRIPT = SCRIPT_DIR / "extract.py"
sys.path.insert(0, str(SCRIPT_DIR))

from extract import extract_to_yaml, extract_with_diagnostics  # noqa: E402


def _bold_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.add_run(text).bold = True
    return paragraph


def _native_list_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_properties.append(level)
    num_properties.append(num_id)
    properties.append(num_properties)
    return paragraph


class ExtractResumeTests(unittest.TestCase):
    def setUp(self):
        self._temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self._temp_dir.cleanup)
        self.temp_dir = Path(self._temp_dir.name)

    def _path(self, name: str = "resume.docx") -> Path:
        return self.temp_dir / name

    @staticmethod
    def _add_identity(document: Document) -> None:
        document.add_paragraph("Jordan Rivers")
        document.add_paragraph("Portland, OR • jordan.rivers@example.com • linkedin.com/in/jordanrivers")

    def _save_valid_shell(self, path: Path) -> Document:
        document = Document()
        self._add_identity(document)
        document.add_heading("Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Software Engineer  2020 – Present | Portland, OR"
        )
        document.add_paragraph("• Built reliable fictional services.")
        document.save(path)
        return document

    def test_extracts_two_employers_in_source_order(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Professional Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Senior Software Engineer  2022 – Present | Portland, OR"
        )
        document.add_paragraph("• Led a fictional platform migration.")
        document.add_paragraph("Contoso Labs")
        document.add_paragraph("Software Engineer")
        document.add_paragraph("2019 — 2022 | Seattle, WA")
        document.add_paragraph("• Improved a fictional deployment pipeline.")
        document.save(path)

        data = extract_to_yaml(str(path))

        self.assertNotIn("employer", data)
        self.assertEqual(
            [employer["company"] for employer in data["employers"]],
            ["Northwind Systems", "Contoso Labs"],
        )
        self.assertEqual(data["employers"][0]["role"], "Senior Software Engineer")
        self.assertEqual(data["employers"][1]["dates"], "2019 — 2022")
        self.assertEqual(
            data["employers"][1]["bullets"],
            ["Improved a fictional deployment pipeline."],
        )

    def test_preserves_promotions_as_repeated_company_entries(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Work History", level=1)
        _bold_paragraph(document, "Fabrikam Systems")
        _bold_paragraph(document, "Senior Platform Engineer")
        document.add_paragraph("2023 – Present | Remote")
        document.add_paragraph("• Led the senior-level fictional work.")
        document.add_paragraph("Platform Engineer  2021 – 2023 | Remote")
        document.add_paragraph("• Delivered the earlier fictional work.")
        document.save(path)

        result = extract_with_diagnostics(path)

        self.assertTrue(result.ok, result.diagnostics)
        assert result.data is not None
        employers = result.data["employers"]
        self.assertEqual(len(employers), 2)
        self.assertEqual([item["company"] for item in employers], ["Fabrikam Systems"] * 2)
        self.assertEqual(
            [item["role"] for item in employers],
            ["Senior Platform Engineer", "Platform Engineer"],
        )

    def test_detects_native_word_list_paragraphs(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Employment", level=1)
        document.add_paragraph(
            "Adventure Works - Site Reliability Engineer  2020 - Present | Austin, TX"
        )
        _native_list_paragraph(document, "Automated a fictional recovery workflow.")
        document.save(path)

        data = extract_to_yaml(str(path))

        self.assertEqual(
            data["employers"][0]["bullets"],
            ["Automated a fictional recovery workflow."],
        )

    def test_keeps_direct_bullets_and_extracts_explicit_projects(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Relevant Experience", level=1)
        document.add_paragraph(
            "Tailspin Toys – Backend Engineer  2021 – Present | Remote"
        )
        document.add_paragraph("• Improved service reliability across a fictional platform.")
        _bold_paragraph(document, "Project Atlas")
        _native_list_paragraph(document, "Built the fictional project ingestion path.")
        _native_list_paragraph(document, "Reduced the fictional project's processing delay.")
        document.save(path)

        data = extract_to_yaml(str(path))
        employer = data["employers"][0]

        self.assertEqual(
            employer["bullets"],
            ["Improved service reliability across a fictional platform."],
        )
        self.assertEqual(
            employer["projects"],
            [{
                "title": "Project Atlas",
                "bullets": [
                    "Built the fictional project ingestion path.",
                    "Reduced the fictional project's processing delay.",
                ],
            }],
        )

    def test_projects_and_experience_sections_preserve_source_order(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Projects", level=1)
        document.add_paragraph(
            "Lakemont University – Student Developer  2022 – 2024 | Lakemont, ST")
        _bold_paragraph(document, "Campus simulator")
        document.add_paragraph("• Built a deterministic fictional route simulator.")
        document.add_heading("Experience", level=1)
        document.add_paragraph(
            "Northwind Systems – Software Engineer  2024 – Present | Remote")
        document.add_paragraph("• Built a reliable fictional backend service.")
        document.save(path)

        data = extract_to_yaml(str(path))
        self.assertEqual(
            [employer["company"] for employer in data["employers"]],
            ["Lakemont University", "Northwind Systems"],
        )

    def test_decorative_drawing_warns_but_image_only_document_fails(self):
        decorated_path = self._path("decorated.docx")
        decorated = Document()
        self._add_identity(decorated)
        decorated.paragraphs[0].add_run()._r.append(OxmlElement("w:drawing"))
        decorated.add_heading("Experience", level=1)
        decorated.add_paragraph(
            "Northwind Systems – Software Engineer  2020 – Present | Remote")
        decorated.add_paragraph("• Built a reliable fictional backend service.")
        decorated.save(decorated_path)

        decorated_result = extract_with_diagnostics(decorated_path)
        self.assertTrue(decorated_result.ok, decorated_result.diagnostics)
        self.assertIn(
            "IGNORED_DECORATIVE_DRAWING",
            [diagnostic.code for diagnostic in decorated_result.diagnostics],
        )

        image_path = self._path("image-only.docx")
        image_only = Document()
        image_only.add_paragraph().add_run()._r.append(OxmlElement("w:drawing"))
        image_only.save(image_path)
        image_result = extract_with_diagnostics(image_path)
        self.assertFalse(image_result.ok)
        self.assertIn(
            "IMAGE_ONLY_DOCUMENT",
            [diagnostic.code for diagnostic in image_result.diagnostics],
        )

    def test_rejects_table_and_multi_column_layouts(self):
        table_path = self._path("table.docx")
        table_document = Document()
        self._add_identity(table_document)
        table_document.add_heading("Experience", level=1)
        table_document.add_table(rows=1, cols=2)
        table_document.save(table_path)

        columns_path = self._path("columns.docx")
        columns_document = Document()
        self._add_identity(columns_document)
        columns_document.add_heading("Experience", level=1)
        columns_document.add_paragraph(
            "Northwind Systems – Engineer  2020 – Present | Portland, OR"
        )
        columns_document.add_paragraph("• Built a fictional service.")
        columns = columns_document.sections[0]._sectPr.xpath("./w:cols")[0]
        columns.set(qn("w:num"), "2")
        columns_document.save(columns_path)

        cases = [
            (table_path, "UNSUPPORTED_TABLE_LAYOUT"),
            (columns_path, "UNSUPPORTED_MULTI_COLUMN_LAYOUT"),
        ]
        for path, code in cases:
            with self.subTest(code=code):
                result = extract_with_diagnostics(path)
                self.assertFalse(result.ok)
                self.assertIsNone(result.data)
                self.assertIn(code, [diagnostic.code for diagnostic in result.diagnostics])

    def test_empty_and_corrupt_files_return_actionable_diagnostics(self):
        empty_path = self._path("empty.docx")
        Document().save(empty_path)
        corrupt_path = self._path("corrupt.docx")
        corrupt_path.write_bytes(b"not a docx package")

        empty = extract_with_diagnostics(empty_path)
        corrupt = extract_with_diagnostics(corrupt_path)

        self.assertEqual(empty.diagnostics[0].code, "EMPTY_DOCUMENT")
        self.assertEqual(corrupt.diagnostics[0].code, "CORRUPT_DOCX")
        process = subprocess.run(
            [sys.executable, str(EXTRACT_SCRIPT), str(corrupt_path)],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(process.returncode, 2)
        self.assertIn("CORRUPT_DOCX", process.stderr)
        self.assertNotIn("Traceback", process.stderr)

    def test_accepts_alternate_heading_pipe_header_and_date_separators(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Employment History", level=1)
        document.add_paragraph(
            "Adventure Works | Site Reliability Engineer | Feb 2018 to Jul 2020 | Austin, TX"
        )
        document.add_paragraph("• Operated fictional production services.")
        document.add_paragraph(
            "Tailspin Toys − Platform Engineer  2020 − Present | Remote"
        )
        document.add_paragraph("• Built a fictional internal platform.")
        document.save(path)

        data = extract_to_yaml(str(path))

        self.assertEqual(len(data["employers"]), 2)
        self.assertEqual(data["employers"][0]["dates"], "Feb 2018 to Jul 2020")
        self.assertEqual(data["employers"][0]["location"], "Austin, TX")
        self.assertEqual(data["employers"][1]["dates"], "2020 − Present")
        self.assertEqual(data["employers"][1]["location"], "Remote")

    def test_ambiguous_experience_fails_instead_of_guessing(self):
        path = self._path()
        document = Document()
        self._add_identity(document)
        document.add_heading("Experience", level=1)
        document.add_paragraph("Northwind Systems")
        document.add_paragraph("This prose has no role or dates.")
        document.save(path)

        result = extract_with_diagnostics(path)

        self.assertFalse(result.ok)
        self.assertIn(
            "AMBIGUOUS_EXPERIENCE_STRUCTURE",
            [diagnostic.code for diagnostic in result.diagnostics],
        )


if __name__ == "__main__":
    unittest.main()

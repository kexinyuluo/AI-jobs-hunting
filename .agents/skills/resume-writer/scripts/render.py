"""Render a tailored resume YAML into DOCX (and optionally PDF).

Reference-based rendering: opens the configured reference DOCX
(config.paths.reference_docx) and replaces content while preserving all original
formatting (fonts, spacing, margins).

Layout: the tailored.yaml and rendered DOCX live in the application folder's
source/ subfolder; the final resume PDF (and the cover-letter PDF) are written to
the application-folder root.

Usage (accepts the app folder, or the source/tailored.yaml path):
    python .agents/skills/resume-writer/scripts/render.py applications/6_drafted/<slug>/
    python .agents/skills/resume-writer/scripts/render.py applications/6_drafted/<slug>/source/tailored.yaml
    python .agents/skills/resume-writer/scripts/render.py applications/6_drafted/<slug>/ --reference path/to/reference.docx
"""

import argparse
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# Self-contained skill: this script lives in the resume-writer skill's scripts/
# folder alongside its _vendor/ copies of the pure toolkit modules. Put both the
# script folder and its _vendor/ on sys.path so sibling scripts (check,
# cover_letter, pdf_convert) and the vendored config/layout/location import.
_HERE = Path(__file__).resolve().parent
for _p in (_HERE, _HERE / "_vendor"):
    if str(_p) not in sys.path and _p.is_dir():
        sys.path.insert(0, str(_p))

import config
from check import application_dir, resume_stem, source_dir, tailored_path

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# RESUME_STEM (and the target-position-aware resume_stem()) live in check.py so
# the renderer and the validator always agree on the output filename.


def prepare_context(data: dict) -> dict:
    """Transform tailored.yaml into render-ready context."""
    name = data.get("name", "")
    contact_line = data.get("contact_line", "")

    summary_bullets = data.get("summary_bullets", [])
    if not summary_bullets and data.get("summary"):
        summary_bullets = [data["summary"]]

    education_line = data.get("education_line", "")

    skills = data.get("skills", [])
    if isinstance(skills, dict):
        skills = [
            {"label": k.replace("_", " ").title(),
             "items": ", ".join(v) if isinstance(v, list) else str(v)}
            for k, v in skills.items() if v
        ]

    employers = data.get("employers", [])
    if not employers and data.get("employer"):
        employers = [data["employer"]]
    if not employers and data.get("experience"):
        employers = []
        for exp in data["experience"]:
            emp = {
                "company": exp.get("company", ""),
                "role": exp.get("role", ""),
                "dates": exp.get("dates", ""),
                "location": exp.get("location", ""),
                "projects": [],
            }
            if exp.get("projects"):
                emp["projects"] = exp["projects"]
            elif exp.get("bullets"):
                emp["projects"] = [{"title": "", "bullets": exp["bullets"]}]
            employers.append(emp)

    return {
        "name": name,
        "contact_line": contact_line,
        "summary_bullets": summary_bullets,
        "education_line": education_line,
        "skills": skills,
        "employers": employers,
    }


# ──────────────────────────────────────────────
# Reference-based rendering (format-preserving)
# ──────────────────────────────────────────────

def _get_text(p_elem):
    """Get concatenated text from a paragraph XML element."""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _split_bold(text):
    """Split '**bold**' markers into (segment, is_bold) tuples."""
    segments = []
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False))
        segments.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False))
    return segments or [(text, False)]


def _set_text(p_elem, text):
    """Replace all content with new text runs, preserving first run's formatting.

    Supports '**bold**' markers for selectively bold segments.
    """
    # Collect first run's formatting before clearing
    first_run = p_elem.find(qn("w:r"))
    if first_run is None:
        # Try inside hyperlinks
        hl = p_elem.find(qn("w:hyperlink"))
        first_run = hl.find(qn("w:r")) if hl is not None else None
    if first_run is None:
        return
    first_rpr = deepcopy(first_run.find(qn("w:rPr")))

    # Remove all runs, hyperlinks, and other inline content (but keep pPr)
    for child in list(p_elem):
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local in ("r", "hyperlink", "bookmarkStart", "bookmarkEnd", "fldSimple"):
            p_elem.remove(child)

    for seg, bold in _split_bold(text):
        r = etree.SubElement(p_elem, qn("w:r"))
        rpr = deepcopy(first_rpr) if first_rpr is not None else None
        if bold:
            if rpr is None:
                rpr = etree.Element(qn("w:rPr"))
            if rpr.find(qn("w:b")) is None:
                rpr.insert(0, etree.Element(qn("w:b")))
        if rpr is not None:
            r.insert(0, rpr)
        t = etree.SubElement(r, qn("w:t"))
        t.text = seg
        t.set(XML_SPACE, "preserve")


def _clone_para(template_elem, text):
    """Deep-copy a paragraph element with new single-run text."""
    new_p = deepcopy(template_elem)
    _set_text(new_p, text)
    return new_p


def _right_tab_pos(body) -> int:
    """Content width in twips (page width - L/R margins) = right-tab position."""
    fallback = 10512  # US Letter, ~0.6in margins
    sect = body.find(qn("w:sectPr"))
    if sect is None:
        return fallback
    pg_sz = sect.find(qn("w:pgSz"))
    pg_mar = sect.find(qn("w:pgMar"))
    try:
        width = float(pg_sz.get(qn("w:w")))
        left = float(pg_mar.get(qn("w:left")))
        right = float(pg_mar.get(qn("w:right")))
        return int(round(width - left - right))
    except (TypeError, ValueError, AttributeError):
        return fallback


def _build_emp_header(template_elem, left, right, tab_pos):
    """Employer header with the dates/location block right-aligned via a tab stop.

    Left text ("Company - Role"), a real tab, then right text
    ("dates | location") that snaps to a right-aligned tab stop at the right
    margin — robust to any text length, unlike space padding.
    """
    new_p = deepcopy(template_elem)
    first_run = new_p.find(qn("w:r"))
    first_rpr = deepcopy(first_run.find(qn("w:rPr"))) if first_run is not None else None

    for child in list(new_p):
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local in ("r", "hyperlink", "bookmarkStart", "bookmarkEnd", "fldSimple"):
            new_p.remove(child)

    ppr = new_p.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.Element(qn("w:pPr"))
        new_p.insert(0, ppr)
    for old_tabs in ppr.findall(qn("w:tabs")):
        ppr.remove(old_tabs)
    tabs = etree.Element(qn("w:tabs"))
    tab = etree.SubElement(tabs, qn("w:tab"))
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(tab_pos))
    ppr.insert(0, tabs)  # w:tabs must precede w:spacing/w:rPr in CT_PPr order

    def _add_run(seg=None, is_tab=False):
        r = etree.SubElement(new_p, qn("w:r"))
        if first_rpr is not None:
            r.insert(0, deepcopy(first_rpr))
        if is_tab:
            etree.SubElement(r, qn("w:tab"))
        else:
            t = etree.SubElement(r, qn("w:t"))
            t.text = seg
            t.set(XML_SPACE, "preserve")

    _add_run(left)
    _add_run(is_tab=True)
    _add_run(right)
    return new_p


def _rebuild_edu(edu_para, edu_line):
    """Rebuild education paragraph: bold 'Education:' + normal content."""
    runs = edu_para.findall(qn("w:r"))
    bold_rpr = deepcopy(runs[0].find(qn("w:rPr"))) if runs else None
    normal_rpr = deepcopy(runs[1].find(qn("w:rPr"))) if len(runs) > 1 else None
    for r in runs:
        edu_para.remove(r)

    r1 = etree.SubElement(edu_para, qn("w:r"))
    if bold_rpr is not None:
        r1.insert(0, bold_rpr)
    t1 = etree.SubElement(r1, qn("w:t"))
    t1.text = "Education:"
    t1.set(XML_SPACE, "preserve")

    r2 = etree.SubElement(edu_para, qn("w:r"))
    if normal_rpr is not None:
        r2.insert(0, normal_rpr)
    t2 = etree.SubElement(r2, qn("w:t"))
    t2.text = f" {edu_line}"
    t2.set(XML_SPACE, "preserve")


def _rebuild_skills(skills_para, skills_data):
    """Rebuild skills paragraph with bold labels and <w:br/> between lines."""
    runs = skills_para.findall(qn("w:r"))
    bold_rpr = deepcopy(runs[0].find(qn("w:rPr"))) if runs else None
    normal_rpr = deepcopy(runs[1].find(qn("w:rPr"))) if len(runs) > 1 else None
    for r in runs:
        skills_para.remove(r)

    for i, skill in enumerate(skills_data):
        r_label = etree.SubElement(skills_para, qn("w:r"))
        if bold_rpr is not None:
            r_label.insert(0, deepcopy(bold_rpr))
        t_label = etree.SubElement(r_label, qn("w:t"))
        t_label.text = f"{skill['label']}:"
        t_label.set(XML_SPACE, "preserve")

        r_val = etree.SubElement(skills_para, qn("w:r"))
        if normal_rpr is not None:
            r_val.insert(0, deepcopy(normal_rpr))
        t_val = etree.SubElement(r_val, qn("w:t"))
        t_val.text = f" {skill['items']}"
        t_val.set(XML_SPACE, "preserve")

        if i < len(skills_data) - 1:
            br = etree.SubElement(r_val, qn("w:br"))
            br.set(qn("w:type"), "textWrapping")


def _find_section(texts, header):
    """Find the index of a section header paragraph."""
    for i, t in enumerate(texts):
        if t.strip() == header:
            return i
    return None


def render_from_reference(ref_path: Path, data: dict, output_path: Path):
    """Render resume by replacing content in the user's reference DOCX."""
    shutil.copy2(ref_path, output_path)
    doc = Document(str(output_path))
    body = doc.element.body
    ctx = prepare_context(data)

    all_p = body.findall(qn("w:p"))
    texts = [_get_text(p) for p in all_p]

    summary_i = _find_section(texts, "Summary")
    edu_i = _find_section(texts, "Education & Skills")
    exp_i = _find_section(texts, "Experience")

    if summary_i is None or edu_i is None or exp_i is None:
        print("Error: could not find all section headers in reference DOCX", file=sys.stderr)
        sys.exit(1)

    # ── Name & Contact ──
    _set_text(all_p[0], ctx["name"])
    _set_text(all_p[1], ctx["contact_line"])

    # ── Summary bullets ──
    old_bullet_idx = [i for i in range(summary_i + 1, edu_i) if texts[i].strip()]
    bullet_template = deepcopy(all_p[old_bullet_idx[0]]) if old_bullet_idx else None
    for idx in reversed(old_bullet_idx):
        body.remove(all_p[idx])
    if bullet_template is not None:
        anchor = all_p[summary_i + 1]  # spacer after "Summary"
        for bt in ctx["summary_bullets"]:
            new_p = _clone_para(bullet_template, bt)
            anchor.addnext(new_p)
            anchor = new_p

    # Re-index
    all_p = body.findall(qn("w:p"))
    texts = [_get_text(p) for p in all_p]
    edu_i = _find_section(texts, "Education & Skills")
    exp_i = _find_section(texts, "Experience")

    # ── Education & Skills ──
    edu_content = [i for i in range(edu_i + 1, exp_i) if texts[i].strip()]
    if edu_content:
        _rebuild_edu(all_p[edu_content[0]], ctx["education_line"])
    if len(edu_content) > 1:
        _rebuild_skills(all_p[edu_content[1]], ctx["skills"])

    # Re-index
    all_p = body.findall(qn("w:p"))
    texts = [_get_text(p) for p in all_p]
    exp_i = _find_section(texts, "Experience")

    # ── Experience ──
    exp_content = [i for i in range(exp_i + 1, len(all_p)) if texts[i].strip()]
    if len(exp_content) < 3:
        doc.save(str(output_path))
        return

    emp_template = deepcopy(all_p[exp_content[0]])
    proj_template = deepcopy(all_p[exp_content[1]])
    bullet_exp_template = deepcopy(all_p[exp_content[2]])

    # Remove all paragraphs after experience spacer
    for idx in reversed(range(exp_i + 2, len(all_p))):
        body.remove(all_p[idx])

    tab_pos = _right_tab_pos(body)
    anchor = all_p[exp_i + 1]  # spacer after "Experience"
    for emp in ctx["employers"]:
        left = f"{emp['company']} – {emp['role']}"
        right = f"{emp['dates']} | {emp['location']}"
        emp_p = _build_emp_header(emp_template, left, right, tab_pos)
        anchor.addnext(emp_p)
        anchor = emp_p

        for proj in emp.get("projects", []):
            if proj.get("title") and proj_template is not None:
                title_p = _clone_para(proj_template, proj["title"])
                anchor.addnext(title_p)
                anchor = title_p
            for b in proj.get("bullets", []):
                bp = _clone_para(bullet_exp_template, b)
                anchor.addnext(bp)
                anchor = bp

    doc.save(str(output_path))


# ──────────────────────────────────────────────
# PDF conversion
# ──────────────────────────────────────────────

def convert_to_pdf(docx_path: Path, output_dir: Path, stem: str) -> Path | None:
    """Convert the rendered resume DOCX to PDF (shared LibreOffice helper)."""
    from pdf_convert import docx_to_pdf
    return docx_to_pdf(docx_path, output_dir, stem)


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Render tailored resume YAML to DOCX + PDF")
    parser.add_argument("yaml_path",
                        help="Application folder, or the source/tailored.yaml path")
    parser.add_argument(
        "--reference", "-r", default=None,
        help="Reference DOCX for format-preserving render "
             "(default: config.paths.reference_docx)")
    parser.add_argument("--no-pdf", action="store_true",
                        help="Skip PDF conversion")
    parser.add_argument("--no-cover-letter", action="store_true",
                        help="Skip rendering the cover letter, even if a .txt exists")
    parser.add_argument("--label", "--position", dest="label", default=None,
                        help="Target-position label appended to the output filename "
                             "(e.g. \"Frontend Engineer\"). Defaults to the tailored.yaml "
                             "`target_position`. Use only when one company needs two "
                             "divergent resumes; leave unset for the standard single resume.")
    parser.add_argument("--skip-checks", action="store_true",
                        help="Skip post-render validation (.agents/skills/resume-writer/scripts/check.py)")
    args = parser.parse_args()

    input_path = Path(args.yaml_path)
    if not input_path.exists():
        print(f"Error: {input_path} not found", file=sys.stderr)
        sys.exit(1)

    app_dir = application_dir(input_path)
    yaml_path = input_path if input_path.is_file() else tailored_path(app_dir)
    if not yaml_path.exists():
        print(f"Error: tailored.yaml not found for {app_dir}", file=sys.stderr)
        sys.exit(1)

    with open(yaml_path) as f:
        data = yaml.safe_load(f)

    if not data:
        print(f"Error: {yaml_path} is empty or invalid", file=sys.stderr)
        sys.exit(1)

    label = args.label if args.label is not None else data.get("target_position", "")
    stem = resume_stem(label)

    # DOCX (an input/intermediate) goes in source/; the final PDF at the root.
    src_dir = source_dir(app_dir)
    src_dir.mkdir(parents=True, exist_ok=True)
    docx_output = src_dir / f"{stem}.docx"

    ref_path = Path(args.reference) if args.reference else config.reference_docx_path()
    if not ref_path.exists():
        print(f"Error: reference {ref_path} not found", file=sys.stderr)
        sys.exit(1)

    # Pre-render one-page estimate (fast, no LibreOffice) so an overflow is
    # visible before the expensive render. Advisory only — check.py remains the
    # authoritative post-render page-count gate.
    try:
        import estimate_layout
        _m = estimate_layout.read_template_metrics(ref_path)
        _p = estimate_layout.derived_params(_m)
        _e = estimate_layout.estimate(data, _p)
        _status, _msg = estimate_layout.verdict(
            _e["total_pt"], _e["budget_pt"], _p["pitch_body"])
        print(f"  Layout: est {_e['total_pt']:.0f}pt / {_e['budget_pt']:.0f}pt "
              f"budget — {_status}: {_msg}")
    except Exception:
        pass

    render_from_reference(ref_path, data, docx_output)
    print(f"  DOCX: {docx_output}")

    pdf_path = None
    if not args.no_pdf:
        pdf_path = convert_to_pdf(docx_output, app_dir, stem)
        if pdf_path:
            print(f"  PDF:  {pdf_path}")
        else:
            print("  PDF:  skipped (install LibreOffice or docx2pdf)")
            print("        Open the DOCX in Word/Google Docs and export as PDF.")

    # ── Cover letters (one per JD/role, from the bundled ..._Application_<role>.txt) ──
    if not args.no_cover_letter:
        import cover_letter
        for role, cl_docx, cl_pdf in cover_letter.render_all_cover_letters(
                app_dir, make_pdf=not args.no_pdf):
            tag = f" [{role}]" if role else ""
            if cl_docx is None:
                print(f"  Cover{tag}: no bundled Application .txt for this role")
                continue
            print(f"  Cover DOCX{tag}: {cl_docx}")
            if cl_pdf:
                print(f"  Cover PDF{tag}:  {cl_pdf}")
            elif not args.no_pdf:
                print(f"  Cover PDF{tag}:  skipped (install LibreOffice or docx2pdf)")

    if not args.skip_checks:
        import check
        print("Validating:")
        if not check.run_checks(yaml_path, pdf_path):
            sys.exit(1)


if __name__ == "__main__":
    main()

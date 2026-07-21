"""Render cover letters from the bundled application .txt files into branded DOCX + PDF.

Cover letters map one-to-one to JDs. For each role in meta.yaml (a ``jobs:`` list, or
the single top-level ``role``) the folder keeps its own copy-paste-friendly bundle:
    applications/<status>/<slug>/<APPLICATION_STEM>_<role>.txt
and it renders to <COVER_STEM>_<role>.{docx,pdf}
(the ``<role>`` slug comes from check.slugify_label, e.g. ``Senior_Platform_Engineer``;
the filename stems come from config).
Each bundle holds three sections, each introduced by a plain title + underline:

    COVER LETTER
    ============
    <Name>
    <City, ST • email • linkedin>           # contact line (has • or @)
    Dear <Company> Hiring Team,             # salutation — NO job-title/subject line
    <body paragraph>                        # one per blank-separated block
    ...
    Sincerely,
    <Name>

    WHY THIS COMPANY & ROLE
    =======================
    <plain-text answer>

    PAST EXPERIENCE
    ===============
    <plain-text answer>

This script extracts the COVER LETTER section, lays it out in the same Arial
branding as the résumé, and writes:
    - source/<COVER_STEM>.docx   (the editable DOCX, alongside the other inputs)
    - <COVER_STEM>.pdf           (the final PDF, at the application-folder root)
render.py calls this automatically after rendering the résumé.

Usage:
    python .agents/skills/resume-writer/scripts/cover_letter.py applications/6_drafted/<slug>/
    python .agents/skills/resume-writer/scripts/cover_letter.py applications/6_drafted/<slug>/ --no-pdf
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Self-contained skill: this script lives in the resume-writer skill's scripts/
# folder alongside its _vendor/ copies of the pure toolkit modules. Put both the
# script folder and its _vendor/ on sys.path so the sibling check module and the
# vendored config/layout/location import regardless of the working directory.
_HERE = Path(__file__).resolve().parent
for _p in (_HERE, _HERE / "_vendor"):
    if str(_p) not in sys.path and _p.is_dir():
        sys.path.insert(0, str(_p))

from check import (APPLICATION_STEM, COVER_STEM, application_dir, application_stem,
                   cover_stem, source_dir)

FONT = "Arial"
NAME_PT = 16
CONTACT_PT = 10
BODY_PT = 11
CONTACT_COLOR = RGBColor(0x55, 0x55, 0x55)

SIGNOFF_RE = re.compile(
    r"^(sincerely|regards|best regards|best|warm regards|kind regards|"
    r"thank you|thanks|respectfully|yours truly)[,.]?$",
    re.I,
)

# Section titles recognized inside the bundled application .txt (case-insensitive).
COVER_TITLES = ("cover letter",)


def parse_bundle(text: str) -> dict:
    """Split the bundled application .txt into {section-title-lower: body}.

    Section headers are a title line immediately followed by an underline made of
    '=' or '-' (setext style), e.g. ``COVER LETTER`` then ``============``.
    """
    lines = text.splitlines()
    sections: dict[str, str] = {}
    title, buf = None, []
    i = 0
    while i < len(lines):
        line = lines[i]
        nxt = lines[i + 1].strip() if i + 1 < len(lines) else ""
        is_rule = len(nxt) >= 3 and set(nxt) <= {"=", "-"}
        if line.strip() and is_rule:
            if title is not None:
                sections[title] = "\n".join(buf).strip("\n")
            title = line.strip().lower()
            buf = []
            i += 2
            continue
        buf.append(line)
        i += 1
    if title is not None:
        sections[title] = "\n".join(buf).strip("\n")
    return sections


def find_application_txt(folder: Path, label: str = "") -> Path | None:
    """Return the bundled application .txt for a folder+label, or None.

    A labeled (per-JD) lookup must match its OWN ``..._Application_<label>.txt``
    exactly — it never falls back to a different role's bundle. Only the unlabeled
    lookup (label="") falls back to any single bundle.
    """
    folder = application_dir(folder)
    exact = folder / f"{application_stem(label)}.txt"
    if exact.exists():
        return exact
    if label:
        return None
    matches = sorted(folder.glob(f"{APPLICATION_STEM}*.txt"))
    return matches[0] if matches else None


def cover_letter_text(folder: Path, label: str = "") -> str | None:
    """The raw cover-letter block: the COVER LETTER section of the bundle."""
    txt = find_application_txt(folder, label)
    if txt is None:
        return None
    raw = txt.read_text()
    sections = parse_bundle(raw)
    for key, body in sections.items():
        if any(t in key for t in COVER_TITLES):
            return body
    return None


def _is_contact(line: str) -> bool:
    return "•" in line or "@" in line or "linkedin.com" in line.lower()


def parse_cover_letter(text: str) -> dict:
    """Split a cover-letter block into header / salutation / body / closing.

    The header keeps only the name (first line) and contact line(s); any other
    header line (e.g. a company/role subject) is dropped so the rendered letter
    starts with the name and info, then goes straight to the salutation.
    """
    lines = [ln.strip() for ln in text.splitlines()]
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()

    sal_idx = next(
        (i for i, l in enumerate(lines) if l.lower().startswith("dear ")), None)
    sign_idx = None
    search_from = (sal_idx + 1) if sal_idx is not None else 0
    for i in range(search_from, len(lines)):
        if SIGNOFF_RE.match(lines[i]):
            sign_idx = i
            break

    header_end = sal_idx if sal_idx is not None else 0
    raw_header = [l for l in lines[:header_end] if l]
    header = [l for i, l in enumerate(raw_header) if i == 0 or _is_contact(l)]
    salutation = lines[sal_idx] if sal_idx is not None else ""

    body_start = (sal_idx + 1) if sal_idx is not None else 0
    body_end = sign_idx if sign_idx is not None else len(lines)
    body_lines = lines[body_start:body_end]

    # Group body into paragraphs by blank lines; join wrapped lines with spaces.
    paragraphs, current = [], []
    for l in body_lines:
        if l:
            current.append(l)
        elif current:
            paragraphs.append(" ".join(current))
            current = []
    if current:
        paragraphs.append(" ".join(current))

    closing = [l for l in lines[sign_idx:] if l] if sign_idx is not None else []

    return {
        "header": header,
        "salutation": salutation,
        "paragraphs": paragraphs,
        "closing": closing,
    }


def _style_run(run, size_pt, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Ensure east-asian text also uses the same face.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def _add_line(doc, text, size_pt, *, bold=False, color=None,
              space_after=0, space_before=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    _style_run(p.add_run(text), size_pt, bold=bold, color=color)
    return p


def build_docx(parsed: dict, docx_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)

    # Header: name first, then contact line(s). No job-title/subject line.
    header = parsed["header"]
    for i, line in enumerate(header):
        last = i == len(header) - 1
        if i == 0:
            _add_line(doc, line, NAME_PT, bold=True, space_after=2)
        else:  # contact line
            _add_line(doc, line, CONTACT_PT, color=CONTACT_COLOR,
                      space_after=2 if not last else 12)

    if parsed["salutation"]:
        _add_line(doc, parsed["salutation"], BODY_PT, space_after=10)

    for para in parsed["paragraphs"]:
        _add_line(doc, para, BODY_PT, space_after=10)

    for i, line in enumerate(parsed["closing"]):
        _add_line(doc, line, BODY_PT,
                  space_before=6 if i == 0 else 0, space_after=0)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return docx_path


def render_cover_letter(folder: Path, label: str = "", make_pdf: bool = True):
    """Render the cover letter for an application folder. Return (docx, pdf|None).

    Reads the COVER LETTER section of the bundled application .txt, writes the
    DOCX into the folder's source/ subfolder, and the final PDF at the folder root.
    Returns (None, None) when no cover-letter text is present.
    """
    app_dir = application_dir(folder)
    body = cover_letter_text(app_dir, label)
    if body is None:
        return None, None

    parsed = parse_cover_letter(body)
    stem = cover_stem(label)
    docx_path = source_dir(app_dir) / f"{stem}.docx"
    build_docx(parsed, docx_path)

    pdf_path = None
    if make_pdf:
        from pdf_convert import docx_to_pdf
        pdf_path = docx_to_pdf(docx_path, app_dir, stem)
    return docx_path, pdf_path


def render_all_cover_letters(folder: Path, make_pdf: bool = True) -> list:
    """Render one cover letter per JD/role — the one-to-one mapping.

    Enumerates roles from meta.yaml (check.application_roles): a ``jobs:`` list
    yields one letter per posting, otherwise the single top-level ``role``. Each
    role's letter comes from its own ``..._Application_<role>.txt`` and renders to
    ``..._Cover_Letter_<role>.{docx,pdf}``. Falls back to a single unlabeled bundle
    for legacy folders whose meta.yaml has no role info.

    Returns a list of (role, docx_path|None, pdf_path|None) — one per role.
    """
    from check import application_roles
    app_dir = application_dir(folder)
    roles = application_roles(app_dir)
    if not roles:
        docx, pdf = render_cover_letter(app_dir, label="", make_pdf=make_pdf)
        return [("", docx, pdf)]
    results = []
    for role in roles:
        docx, pdf = render_cover_letter(app_dir, label=role, make_pdf=make_pdf)
        results.append((role, docx, pdf))
    return results


def main():
    parser = argparse.ArgumentParser(
        description="Render one cover letter per JD from the bundled application .txt files")
    parser.add_argument("path",
                        help="Application folder (or a bundled Application .txt file)")
    parser.add_argument("--label", "--position", dest="label", default=None,
                        help="Render only this one role's ..._Application_<label>.txt "
                             "(default: render every role in meta.yaml, one per JD)")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF conversion")
    args = parser.parse_args()

    folder = application_dir(Path(args.path))

    from pdf_convert import PdfConversionError
    try:
        if args.label is not None:
            if find_application_txt(folder, args.label) is None:
                print(f"Error: no {application_stem(args.label)}.txt found under {folder}",
                      file=sys.stderr)
                sys.exit(1)
            docx_path, pdf_path = render_cover_letter(
                folder, label=args.label, make_pdf=not args.no_pdf)
            results = [(args.label, docx_path, pdf_path)]
        else:
            results = render_all_cover_letters(folder, make_pdf=not args.no_pdf)
            if not any(docx is not None for _, docx, _ in results):
                print(f"Error: no bundled application .txt found under {folder}",
                      file=sys.stderr)
                sys.exit(1)
    except PdfConversionError as exc:
        print(f"Error: cover-letter PDF conversion failed: {exc}", file=sys.stderr)
        sys.exit(1)

    for role, docx_path, pdf_path in results:
        tag = f" [{role}]" if role else ""
        if docx_path is None:
            print(f"  Cover{tag}: no bundle found (expected {application_stem(role)}.txt)")
            continue
        print(f"  Cover DOCX{tag}: {docx_path}")
        if pdf_path:
            print(f"  Cover PDF{tag}:  {pdf_path}")
        elif not args.no_pdf:
            print(f"  Cover PDF{tag}:  skipped (install LibreOffice or docx2pdf)")
            print("              Open the DOCX in Word/Google Docs and export as PDF.")


if __name__ == "__main__":
    main()

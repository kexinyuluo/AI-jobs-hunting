"""Extract a standard single-column DOCX resume into tailored YAML.

The parser deliberately rejects layouts it cannot interpret reliably. Call
``extract_with_diagnostics`` when diagnostics are needed; ``extract_to_yaml``
retains the original success-path API and raises ``ExtractionError`` on fatal
input.
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn


SECTION_KEYWORDS = {
    "summary": {
        "summary", "professional summary", "career summary", "objective",
        "profile", "professional profile", "about",
    },
    "experience": {
        "experience", "work experience", "professional experience",
        "relevant experience", "employment", "employment history",
        "work history", "career history", "professional background",
        "internship experience", "internships", "consulting experience",
        "contract experience",
    },
    "education_skills": {"education & skills", "education and skills"},
    "skills": {
        "skills", "technical skills", "core competencies", "technologies",
        "technical competencies",
    },
    "education": {"education", "academic background", "academics"},
    "certifications": {
        "certifications", "certificates", "licenses",
        "certifications & licenses", "certifications and licenses",
    },
    "projects": {"projects", "personal projects", "side projects"},
}

_BULLET_RE = re.compile(r"^(?:[\u2022\u2023\u25e6\u2043\u25aa\u25cf]\s*|[-*>]\s+)")
_DASH_CHARS = r"\-\u2010\u2011\u2012\u2013\u2014\u2212"
_COMPANY_ROLE_RE = re.compile(rf"\s+[{_DASH_CHARS}]\s+")
_MONTH = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|"
    r"Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|"
    r"Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.?"
)
_DATE_ENDPOINT = rf"(?:{_MONTH}\s+\d{{4}}|\d{{1,2}}/\d{{4}}|\d{{4}}|Present|Current|Now)"
_DATE_RANGE_RE = re.compile(
    rf"(?P<dates>{_DATE_ENDPOINT}\s*(?:[{_DASH_CHARS}]|to)\s*{_DATE_ENDPOINT})",
    re.IGNORECASE,
)
_LOCATION_RE = re.compile(
    r"(?:\bremote\b|\bhybrid\b|\bon[- ]?site\b|,\s*[A-Z]{2}\b|,\s*[A-Za-z][A-Za-z .'-]+$)",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class Diagnostic:
    """An actionable extraction problem."""

    severity: str
    code: str
    message: str


@dataclass
class ExtractionResult:
    """Structured extraction output plus diagnostics."""

    data: dict[str, Any] | None
    diagnostics: list[Diagnostic]

    @property
    def ok(self) -> bool:
        return self.data is not None and not any(d.severity == "error" for d in self.diagnostics)

    @property
    def success(self) -> bool:
        return self.ok

    @property
    def fatal(self) -> bool:
        return not self.ok


class ExtractionError(ValueError):
    """Raised by the compatibility API when extraction is unsafe."""

    def __init__(self, diagnostics: list[Diagnostic]):
        self.diagnostics = diagnostics
        message = "; ".join(d.message for d in diagnostics if d.severity == "error")
        super().__init__(message or "DOCX extraction failed")


@dataclass(frozen=True)
class _Paragraph:
    text: str
    style: str
    bold: bool
    native_list: bool

    @property
    def bullet(self) -> bool:
        return self.native_list or looks_like_bullet(self.text)


def _normalized_heading(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text.strip().lower().rstrip(":"))
    return normalized.replace("＆", "&")


def classify_section(text: str) -> str | None:
    normalized = _normalized_heading(text)
    for section, keywords in SECTION_KEYWORDS.items():
        if normalized in keywords:
            return section
    return None


def _num_pr_is_active(num_pr: Any) -> bool:
    if num_pr is None:
        return False
    num_id = getattr(num_pr, "numId", None)
    if num_id is None:
        return True
    value = num_id.get(qn("w:val"))
    return value != "0"


def _has_native_list(paragraph: Any) -> bool:
    """Detect direct or style-inherited Word numbering (``w:numPr``)."""
    p_pr = paragraph._p.pPr
    if p_pr is not None and _num_pr_is_active(p_pr.numPr):
        return True

    style = paragraph.style
    seen: set[int] = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        style_p_pr = style._element.pPr
        if style_p_pr is not None and _num_pr_is_active(style_p_pr.numPr):
            return True
        style = style.base_style
    return False


def _paragraphs(document: DocumentType) -> list[_Paragraph]:
    output = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        output.append(_Paragraph(
            text=text,
            style=paragraph.style.name if paragraph.style else "Normal",
            bold=any(run.bold is True for run in paragraph.runs),
            native_list=_has_native_list(paragraph),
        ))
    return output


def extract_paragraphs(docx_path: str) -> list[dict[str, Any]]:
    """Compatibility helper returning paragraph metadata for a valid DOCX."""
    document = Document(docx_path)
    return [
        {
            "text": paragraph.text,
            "style": paragraph.style,
            "bold": paragraph.bold,
            "native_list": paragraph.native_list,
            "bullet": paragraph.bullet,
        }
        for paragraph in _paragraphs(document)
    ]


def is_heading(para: dict[str, Any] | _Paragraph) -> bool:
    text = para.text if isinstance(para, _Paragraph) else para["text"]
    style = para.style if isinstance(para, _Paragraph) else para["style"]
    bold = para.bold if isinstance(para, _Paragraph) else para["bold"]
    return (
        "heading" in style.lower()
        or (bold and len(text) < 60)
        or (text.isupper() and len(text) < 60)
    )


def looks_like_bullet(text: str) -> bool:
    return bool(_BULLET_RE.match(text))


def clean_bullet(text: str) -> str:
    return _BULLET_RE.sub("", text, count=1).strip()


def _clean_paragraph_bullet(paragraph: _Paragraph) -> str:
    return clean_bullet(paragraph.text) if looks_like_bullet(paragraph.text) else paragraph.text.strip()


def _layout_diagnostics(document: DocumentType) -> list[Diagnostic]:
    diagnostics = []
    if document.tables:
        diagnostics.append(Diagnostic(
            "error",
            "UNSUPPORTED_TABLE_LAYOUT",
            "The DOCX contains a table. Convert the resume to ordinary single-column paragraphs before extracting.",
        ))

    for index, section in enumerate(document.sections, start=1):
        columns = section._sectPr.findall(qn("w:cols"))
        if not columns:
            continue
        num = columns[0].get(qn("w:num"))
        explicit_columns = columns[0].findall(qn("w:col"))
        if (num and int(num) > 1) or len(explicit_columns) > 1:
            diagnostics.append(Diagnostic(
                "error",
                "UNSUPPORTED_MULTI_COLUMN_LAYOUT",
                f"Section {index} uses multiple columns. Convert it to one column before extracting.",
            ))

    if next(document.element.iter(qn("w:txbxContent")), None) is not None:
        diagnostics.append(Diagnostic(
            "error",
            "UNSUPPORTED_TEXT_BOX",
            "The DOCX contains text-box content, whose reading order is ambiguous. Convert it to ordinary paragraphs.",
        ))
    if _has_drawings(document):
        diagnostics.append(Diagnostic(
            "warning",
            "IGNORED_DECORATIVE_DRAWING",
            "The DOCX contains a drawing or image; extraction ignores it and "
            "uses the surrounding paragraph text.",
        ))
    return diagnostics


def _has_drawings(document: DocumentType) -> bool:
    return (
        next(document.element.iter(qn("w:drawing")), None) is not None
        or next(document.element.iter(qn("w:pict")), None) is not None
    )


def _split_sections(
    paragraphs: list[_Paragraph],
) -> tuple[
    list[_Paragraph],
    dict[str, list[_Paragraph]],
    list[tuple[str, list[_Paragraph]]],
]:
    preamble = []
    sections: dict[str, list[_Paragraph]] = {}
    ordered_blocks: list[tuple[str, list[_Paragraph]]] = []
    current = None
    current_block: list[_Paragraph] | None = None
    for paragraph in paragraphs:
        section = classify_section(paragraph.text)
        if section and (is_heading(paragraph) or len(paragraph.text) < 60):
            current = section
            sections.setdefault(section, [])
            current_block = []
            ordered_blocks.append((section, current_block))
        elif current is None:
            preamble.append(paragraph)
        else:
            sections.setdefault(current, []).append(paragraph)
            assert current_block is not None
            current_block.append(paragraph)
    return preamble, sections, ordered_blocks


def _extract_role_dates_location(text: str) -> tuple[str, str, str] | None:
    match = _DATE_RANGE_RE.search(text)
    if not match:
        return None
    role = text[:match.start()].strip(" |\t")
    location = text[match.end():].strip(" |\t")
    if not role:
        return None
    return role, match.group("dates").strip(), location


def _extract_dates_location(text: str) -> tuple[str, str] | None:
    match = _DATE_RANGE_RE.search(text)
    if not match or text[:match.start()].strip(" |\t"):
        return None
    return match.group("dates").strip(), text[match.end():].strip(" |\t")


def _parse_combined_header(text: str) -> dict[str, str] | None:
    """Parse ``Company – Role  Dates | Location`` and pipe variants."""
    pipe_parts = [part.strip() for part in text.split("|")]
    if len(pipe_parts) >= 3:
        date_index = next((i for i, part in enumerate(pipe_parts) if _DATE_RANGE_RE.fullmatch(part)), None)
        if date_index is not None and date_index >= 2:
            company = pipe_parts[0]
            role = " | ".join(pipe_parts[1:date_index])
            location = " | ".join(pipe_parts[date_index + 1:])
            if company and role:
                return {
                    "company": company,
                    "role": role,
                    "dates": pipe_parts[date_index],
                    "location": location,
                }

    match = _DATE_RANGE_RE.search(text)
    if not match:
        return None
    prefix = text[:match.start()].strip(" |\t")
    location = text[match.end():].strip(" |\t")
    company_role = _COMPANY_ROLE_RE.split(prefix, maxsplit=1)
    if len(company_role) != 2 or not all(part.strip() for part in company_role):
        return None
    return {
        "company": company_role[0].strip(),
        "role": company_role[1].strip(),
        "dates": match.group("dates").strip(),
        "location": location,
    }


def _without_dates_company_role(text: str) -> tuple[str, str] | None:
    if _DATE_RANGE_RE.search(text):
        return None
    parts = _COMPANY_ROLE_RE.split(text.strip(), maxsplit=1)
    if len(parts) == 2 and all(parts):
        return parts[0].strip(), parts[1].strip()
    return None


def _short_header_line(paragraph: _Paragraph) -> bool:
    return not paragraph.bullet and len(paragraph.text) <= 100 and "\n" not in paragraph.text


def _looks_like_location(text: str) -> bool:
    return len(text) <= 60 and bool(_LOCATION_RE.search(text))


def _consume_employer_header(
    paragraphs: list[_Paragraph],
    index: int,
    current_company: str | None,
) -> tuple[dict[str, str], int] | None:
    paragraph = paragraphs[index]
    if paragraph.bullet:
        return None

    combined = _parse_combined_header(paragraph.text)
    if combined:
        return combined, 1

    role_dates = _extract_role_dates_location(paragraph.text)
    if role_dates and current_company:
        role, dates, location = role_dates
        return {
            "company": current_company,
            "role": role,
            "dates": dates,
            "location": location,
        }, 1

    if index + 1 >= len(paragraphs) or not _short_header_line(paragraph):
        return None

    next_paragraph = paragraphs[index + 1]
    company_role = _without_dates_company_role(paragraph.text)
    dates_location = _extract_dates_location(next_paragraph.text)
    if company_role and dates_location:
        company, role = company_role
        dates, location = dates_location
        return {
            "company": company,
            "role": role,
            "dates": dates,
            "location": location,
        }, 2

    if current_company and dates_location:
        dates, location = dates_location
        return {
            "company": current_company,
            "role": paragraph.text,
            "dates": dates,
            "location": location,
        }, 2

    next_role_dates = _extract_role_dates_location(next_paragraph.text)
    if next_role_dates:
        role, dates, location = next_role_dates
        return {
            "company": paragraph.text,
            "role": role,
            "dates": dates,
            "location": location,
        }, 2

    if index + 2 >= len(paragraphs) or not _short_header_line(next_paragraph):
        return None
    third_paragraph = paragraphs[index + 2]
    dates_location = _extract_dates_location(third_paragraph.text)
    if not dates_location:
        return None

    company = paragraph.text
    role = next_paragraph.text
    consumed = 3

    dates, location = dates_location
    if not location and index + consumed < len(paragraphs):
        location_paragraph = paragraphs[index + consumed]
        if not location_paragraph.bullet and _looks_like_location(location_paragraph.text):
            location = location_paragraph.text
            consumed += 1
    return {
        "company": company,
        "role": role,
        "dates": dates,
        "location": location,
    }, consumed


def _is_project_header(paragraphs: list[_Paragraph], index: int) -> bool:
    paragraph = paragraphs[index]
    text = paragraph.text.rstrip(":").strip()
    if paragraph.bullet or not text or len(text) > 95 or _DATE_RANGE_RE.search(text):
        return False
    if index + 1 >= len(paragraphs) or not paragraphs[index + 1].bullet:
        return False
    explicit_label = bool(re.match(r"^(?:project|initiative)\b", text, re.IGNORECASE))
    return explicit_label or paragraph.bold or "heading" in paragraph.style.lower()


def _parse_experience(paragraphs: list[_Paragraph]) -> tuple[list[dict[str, Any]], list[Diagnostic]]:
    employers: list[dict[str, Any]] = []
    diagnostics = []
    current_employer: dict[str, Any] | None = None
    current_project: dict[str, Any] | None = None
    index = 0

    while index < len(paragraphs):
        paragraph = paragraphs[index]
        header = _consume_employer_header(
            paragraphs,
            index,
            current_employer["company"] if current_employer else None,
        )
        if header:
            current_employer, consumed = header
            employers.append(current_employer)
            current_project = None
            index += consumed
            continue

        if paragraph.bullet:
            if current_employer is None:
                diagnostics.append(Diagnostic(
                    "error",
                    "AMBIGUOUS_EXPERIENCE_STRUCTURE",
                    f"Achievement bullet appears before an employer header: {paragraph.text!r}.",
                ))
                index += 1
                continue
            bullet = _clean_paragraph_bullet(paragraph)
            if current_project is not None:
                current_project["bullets"].append(bullet)
            else:
                current_employer.setdefault("bullets", []).append(bullet)
            index += 1
            continue

        if current_employer is not None and _is_project_header(paragraphs, index):
            current_project = {"title": paragraph.text.rstrip(":").strip(), "bullets": []}
            current_employer.setdefault("projects", []).append(current_project)
            index += 1
            continue

        diagnostics.append(Diagnostic(
            "error",
            "AMBIGUOUS_EXPERIENCE_STRUCTURE",
            f"Could not identify this Experience line as an employer header, project header, or bullet: {paragraph.text!r}.",
        ))
        index += 1

    if not employers:
        diagnostics.append(Diagnostic(
            "error",
            "MISSING_EMPLOYER_HEADERS",
            "No complete employer headers were found. Use 'Company – Role  Dates | Location' or separate company/role/date lines.",
        ))
    for employer in employers:
        if not employer.get("bullets") and not employer.get("projects"):
            diagnostics.append(Diagnostic(
                "error",
                "EMPTY_EMPLOYER_ENTRY",
                f"No achievement bullets were found for {employer['company']} / {employer['role']}.",
            ))
    return employers, diagnostics


def _base_result(preamble: list[_Paragraph]) -> dict[str, Any]:
    name = ""
    contact = ""
    for paragraph in preamble:
        if not name and len(paragraph.text) < 60:
            name = paragraph.text
            continue
        lowered = paragraph.text.lower()
        if not contact and ("@" in paragraph.text or "linkedin" in lowered or "example.com" in lowered):
            contact = paragraph.text
    return {
        "name": name,
        "contact_line": contact,
        "summary_bullets": [],
        "education_line": "",
        "skills": [],
        "employers": [],
    }


def _populate_non_experience_sections(
    result: dict[str, Any],
    sections: dict[str, list[_Paragraph]],
) -> None:
    result["summary_bullets"] = [
        _clean_paragraph_bullet(paragraph) for paragraph in sections.get("summary", [])
    ]

    combined = sections.get("education_skills", [])
    if combined:
        lines = []
        for paragraph in combined:
            lines.extend(line.strip() for line in paragraph.text.splitlines() if line.strip())
        for line in lines:
            cleaned = clean_bullet(line) if looks_like_bullet(line) else line
            if ":" in cleaned:
                label, items = (part.strip() for part in cleaned.split(":", 1))
                if label.lower() == "education":
                    result["education_line"] = items
                else:
                    result["skills"].append({"label": label, "items": items})
            elif not result["education_line"]:
                result["education_line"] = cleaned
        return

    if sections.get("education"):
        result["education_line"] = " ".join(
            _clean_paragraph_bullet(paragraph) for paragraph in sections["education"]
        )
    for paragraph in sections.get("skills", []):
        cleaned = _clean_paragraph_bullet(paragraph)
        if ":" in cleaned:
            label, items = (part.strip() for part in cleaned.split(":", 1))
        else:
            label, items = "Skills", cleaned
        result["skills"].append({"label": label, "items": items})


def extract_with_diagnostics(docx_path: str | Path) -> ExtractionResult:
    """Extract a DOCX, returning fatal errors instead of partial silent output."""
    path = Path(docx_path)
    if not path.is_file():
        return ExtractionResult(None, [Diagnostic(
            "error",
            "FILE_NOT_FOUND",
            f"DOCX file not found: {path}",
        )])

    try:
        document = Document(str(path))
    except Exception:  # python-docx exposes several package/XML exception types
        return ExtractionResult(None, [Diagnostic(
            "error",
            "CORRUPT_DOCX",
            "Could not open resume document as a valid DOCX package. "
            "Re-save or replace the file.",
        )])

    diagnostics = _layout_diagnostics(document)
    if any(item.severity == "error" for item in diagnostics):
        return ExtractionResult(None, diagnostics)
    paragraphs = _paragraphs(document)
    if not paragraphs:
        diagnostics.append(Diagnostic(
            "error",
            "IMAGE_ONLY_DOCUMENT" if _has_drawings(document) else "EMPTY_DOCUMENT",
            (
                "Resume content is image-only and cannot be extracted as text. "
                "Provide a paragraph-based DOCX."
                if _has_drawings(document)
                else "Resume document contains no extractable paragraph text."
            ),
        ))
        return ExtractionResult(None, diagnostics)
    preamble, sections, ordered_blocks = _split_sections(paragraphs)
    if "experience" not in sections:
        return ExtractionResult(None, [Diagnostic(
            "error",
            "MISSING_EXPERIENCE_SECTION",
            "No recognized Experience heading was found. Add a heading such as 'Experience', 'Work History', or 'Employment History'.",
        )])
    if not sections["experience"]:
        return ExtractionResult(None, [Diagnostic(
            "error",
            "EMPTY_EXPERIENCE_SECTION",
            "The Experience section contains no paragraph content.",
        )])

    result = _base_result(preamble)
    _populate_non_experience_sections(result, sections)
    # Preserve the source order when Projects and Experience appear as separate
    # sections. A project block still needs an employer-style owner header; the
    # regular ambiguity diagnostics fail closed when ownership is unclear.
    experience_paragraphs = [
        paragraph
        for section, block in ordered_blocks
        if section in {"experience", "projects"}
        for paragraph in block
    ]
    employers, experience_diagnostics = _parse_experience(experience_paragraphs)
    diagnostics.extend(experience_diagnostics)
    if any(item.severity == "error" for item in diagnostics):
        return ExtractionResult(None, diagnostics)
    result["employers"] = employers
    return ExtractionResult(result, diagnostics)


def extract_to_yaml(docx_path: str) -> dict[str, Any]:
    """Compatibility success-path API; raises ``ExtractionError`` on fatal input."""
    result = extract_with_diagnostics(docx_path)
    if not result.ok:
        raise ExtractionError(result.diagnostics)
    assert result.data is not None
    return result.data


def _print_diagnostics(path: Path, diagnostics: list[Diagnostic]) -> None:
    print(f"Error: unable to safely extract {path}", file=sys.stderr)
    for diagnostic in diagnostics:
        print(f"  - [{diagnostic.code}] {diagnostic.message}", file=sys.stderr)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Extract resume content from DOCX to YAML")
    parser.add_argument("docx_path", help="Path to the DOCX resume file")
    parser.add_argument(
        "--output", "-o", default=None,
        help="Output YAML path (default: prints to stdout)",
    )
    args = parser.parse_args(argv)
    docx_path = Path(args.docx_path)

    extraction = extract_with_diagnostics(docx_path)
    if not extraction.ok:
        _print_diagnostics(docx_path, extraction.diagnostics)
        return 2
    assert extraction.data is not None

    yaml_str = yaml.dump(
        extraction.data,
        default_flow_style=False,
        allow_unicode=True,
        sort_keys=False,
        width=120,
    )
    header = (
        f"# Extracted from: {docx_path.name}\n"
        "# Review and clean up the extracted content.\n"
        "# Schema: name, contact_line, summary_bullets, education_line, skills, employers\n\n"
    )
    output = header + yaml_str

    if args.output:
        output_path = Path(args.output)
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(output, encoding="utf-8")
        except OSError as exc:
            print(f"Error: could not write {output_path}: {exc}", file=sys.stderr)
            return 2
        print(f"Extracted resume content to {output_path}")
    else:
        print(output)
    print("Please review the output and fix any extraction errors.", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

"""Drive the tracked public resume fixture matrix end to end."""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPTS = Path(__file__).resolve().parents[1]
REPO_ROOT = Path(__file__).resolve().parents[4]
FIXTURES = REPO_ROOT / "examples" / "fixtures" / "resume-writer"
EXTRACT_SCRIPT = SCRIPTS / "extract.py"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from extract import extract_with_diagnostics  # noqa: E402
from resume_schema import normalize_resume  # noqa: E402


SUPPORTED = (
    "legacy-project-focused",
    "chronological-two-employer",
    "same-company-promotion",
    "hybrid-role-bullets-projects",
    "new-grad-internships-projects",
    "concurrent-contractor-roles",
)


def _load(path: Path):
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def _native_bullet(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.extend((ilvl, num_id))
    p_pr.append(num_pr)


def build_recipe(recipe: dict, output: Path):
    """Materialize the small public ``docx-*-v1`` recipes."""
    document = Document()
    spec = recipe.get("document") or {}

    columns = int(spec.get("columns") or 1)
    if columns > 1:
        sect_pr = document.sections[0]._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        cols.set(qn("w:num"), str(columns))

    if spec.get("tables"):
        table = document.add_table(rows=1, cols=2)
        fragments = recipe.get("synthetic_fragments") or {}
        table.cell(0, 0).text = "\n".join(fragments.get("left_column") or ["Synthetic left"])
        table.cell(0, 1).text = "\n".join(fragments.get("right_column") or ["Synthetic right"])

    for item in recipe.get("paragraphs") or []:
        style = item.get("style") or "Normal"
        try:
            paragraph = document.add_paragraph(style=style)
        except KeyError:
            paragraph = document.add_paragraph(style="Normal")
        run = paragraph.add_run(str(item.get("text") or ""))
        run.bold = bool(item.get("bold"))
        if item.get("list") == "bullet":
            _native_bullet(paragraph)
    document.save(output)


class PublicFixtureMatrixTests(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())
        self.addCleanup(shutil.rmtree, self.tmp, ignore_errors=True)

    def test_index_declares_every_scenario(self):
        index = _load(FIXTURES / "index.yaml")
        ids = {row["id"] for row in index["scenarios"]}
        for scenario in SUPPORTED + (
                "unsupported-two-column", "empty-corrupt",
                "_test_application_multi-experience"):
            self.assertIn(scenario, ids)
        self.assertTrue(index["public_safe"])

    def test_supported_scenarios_match_expected_canonical_yaml(self):
        for scenario in SUPPORTED:
            with self.subTest(scenario=scenario):
                root = FIXTURES / scenario
                recipe = _load(root / "input.yaml")
                expected = normalize_resume(_load(root / "expected-canonical.yaml"))
                if recipe["kind"] == "resume-yaml":
                    actual = normalize_resume(recipe["resume"])
                    self.assertEqual(actual, expected)
                    recipe = _load(root / "input-docx.yaml")

                committed = root / "input.docx"
                self.assertTrue(committed.is_file(), committed)
                generated = self.tmp / f"{scenario}.docx"
                build_recipe(recipe, generated)
                for docx in (committed, generated):
                    result = extract_with_diagnostics(docx)
                    self.assertTrue(
                        result.ok,
                        [(item.code, item.message) for item in result.diagnostics],
                    )
                    actual = normalize_resume(result.data)
                    self.assertEqual(actual, expected)
                    self.assertEqual(
                        [employer["company"] for employer in actual["employers"]],
                        [employer["company"] for employer in expected["employers"]],
                    )

    def test_unsupported_two_column_matches_expected_diagnostic(self):
        root = FIXTURES / "unsupported-two-column"
        recipe = _load(root / "input.yaml")
        expected = _load(root / "expected-diagnostics.yaml")
        docx = root / "input.docx"
        self.assertTrue(docx.is_file(), docx)
        result = extract_with_diagnostics(docx)
        self.assertFalse(result.ok)
        self.assertIsNone(result.data)
        self.assertEqual(
            [(item.code, item.message) for item in result.diagnostics],
            [(item["code"], item["message"]) for item in expected["diagnostics"]],
        )
        self._assert_cli_failure(docx, expected["diagnostics"])

    def test_empty_and_corrupt_cases_match_expected_diagnostics(self):
        root = FIXTURES / "empty-corrupt"
        fixture = _load(root / "input.yaml")
        expected = {
            row["id"]: row for row in _load(root / "expected-diagnostics.yaml")["cases"]
        }
        for case in fixture["cases"]:
            with self.subTest(case=case["id"]):
                path = root / f"{case['id']}.docx"
                self.assertTrue(path.is_file(), path)
                result = extract_with_diagnostics(path)
                self.assertFalse(result.ok)
                self.assertIsNone(result.data)
                self.assertEqual(
                    [(item.code, item.message) for item in result.diagnostics],
                    [
                        (item["code"], item["message"])
                        for item in expected[case["id"]]["diagnostics"]
                    ],
                )
                self._assert_cli_failure(
                    path, expected[case["id"]]["diagnostics"])

    def _assert_cli_failure(self, path: Path, expected: list[dict]):
        process = subprocess.run(
            [sys.executable, str(EXTRACT_SCRIPT), str(path)],
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(process.returncode, 2)
        self.assertEqual(process.stdout, "")
        self.assertNotIn("Traceback", process.stderr)
        for item in expected:
            self.assertIn(f"[{item['code']}]", process.stderr)
            self.assertIn(item["message"], process.stderr)


if __name__ == "__main__":
    unittest.main()
